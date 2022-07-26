# -*- coding: utf-8 -*-
import json
import pyperclip
import docx
import pygame
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches


def out_docx(sorted_dict):
    LIGHT_GRAY = [parse_xml(r'<w:shd {} w:fill="DDDDDD"/>'.format(nsdecls('w')))]
    DARK_GRAY = [parse_xml(r'<w:shd {} w:fill="888888"/>'.format(nsdecls('w')))]
    height = 0
    width = 0
    LIGHT_GRAY_COUNT = 0
    DARK_GRAY_COUNT = 0
    doc = docx.Document()
    doc.add_heading('RESULT', 0)
    data = []
    action_count = 0
    for i in sorted_dict:
        action_count += 1
        data.append([action_count, i, sorted_dict[i]])
    if len(data) == 0:
        data = [['1', 'None', '--']]
    table = doc.add_table(rows=2, cols=len(data) + 2)
    row = table.rows[0].cells
    row[0].text = '№'
    row[1].text = 'Игрок'
    row[2].text = 'Баллы'
    for num, name, points in data:
        row = table.add_row().cells
        row[0].text = str(num)
        row[1].text = str(name)
        row[2].text = str(points)
    for i in range(0, len(data) * 3 + 6):
        if i % 2 == 0:
            LIGHT_GRAY.append(parse_xml(r'<w:shd {} w:fill="DDDDDD"/>'.format(nsdecls('w'))))
            cell = table.cell(width, height)
            cell._tc.get_or_add_tcPr().append(LIGHT_GRAY[LIGHT_GRAY_COUNT])
            LIGHT_GRAY_COUNT += 1
            height += 1

            if height == 3:
                height = 0
                width += 1
        else:
            DARK_GRAY.append(parse_xml(r'<w:shd {} w:fill="888888"/>'.format(nsdecls('w'))))
            cell = table.cell(width, height)
            cell._tc.get_or_add_tcPr().append(DARK_GRAY[DARK_GRAY_COUNT])
            DARK_GRAY_COUNT += 1
            height += 1

            if height == 3:
                height = 0
                width += 1
    for cell in table.columns[0].cells:
        cell.width = Inches(0.2)
    for cell in table.columns[1].cells:
        cell.width = Inches(10)
    for cell in table.columns[2].cells:
        cell.width = Inches(0.5)

    doc.save('documents/table.docx')


def write(points, name_club, NewPoints):
    with open('data.json', 'r') as j:
        content = json.loads(j.read())
    if points is not None:
        content[name_club] = points + NewPoints
        with open('data.json', 'w') as json_file:
            json.dump(content, json_file, indent=2, sort_keys=True, ensure_ascii=False)
    else:
        content[name_club] = NewPoints
        with open('data.json', 'w') as json_file:
            json.dump(content, json_file, indent=2, sort_keys=True, ensure_ascii=False)


def finder(name_club):
    with open('data.json', 'r') as j:
        content = json.loads(j.read())
        try:
            a = content[name_club]
            return a
        except:
            return None


def top():
    with open('data.json', 'r') as j:  # нет бага
        content = json.loads(j.read())
    myDictionary = content
    sorted_dict = {}

    while len(myDictionary) != 0:
        BiggerNumberWithKey = {None: 0}
        for i in myDictionary:
            for k in BiggerNumberWithKey:
                if myDictionary[i] > BiggerNumberWithKey[k]:
                    BiggerNumberWithKey = {i: myDictionary[i]}
                elif myDictionary == BiggerNumberWithKey[k]:
                    BiggerNumberWithKey[i] = myDictionary[i]
                    break
        for i in BiggerNumberWithKey:
            del myDictionary[i]
            sorted_dict[i] = BiggerNumberWithKey[i]
    return sorted_dict


def chart(tops):
    a = ''
    NumberOfChart = 0
    for clubs in tops:
        NumberOfChart += 1
        a += f'№{NumberOfChart}| {clubs} - {tops[clubs]}\n'
        pyperclip.copy(a)
    if NumberOfChart == 0:
        pyperclip.copy('empty')


def clear():
    with open('data.json', 'w') as json_file:
        json.dump({}, json_file, indent=2, sort_keys=True, ensure_ascii=False)


thickarrow_strings = (
    "                        ",
    "     XXXXXX  XXXXXX     ",
    "    X      XX      X    ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "           XX           ",
    "    X      XX      X    ",
    "     XXXXXX  XXXXXX     ",
    "                        ")

DARKGRAY = (50, 50, 50)
LIGHTGRAY = (150, 150, 150)
SUPERLIGHTGRAY = (220, 220, 220)
LIGHTBLUE = (34, 113, 179)
